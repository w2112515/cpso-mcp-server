"""
Ingestion Node for CPSO-Protocol Multi-Agent System
Handles parsing and ingestion of various document formats including markdown, PDF, and DOCX files.
Also includes the legacy Intelligence Officer functionality for backward compatibility.
"""

import os
import logging
from typing import List, Dict, Any, Optional
from ..state.schema import GlobalState, GlobalStateStatus, RawIntelligence, InputAttachment
from ..utils.llm_router import get_llm
from ..utils.artifact import ArtifactManager

# Set up logger
logger = logging.getLogger(__name__)

# Conditional imports for optional dependencies
PDF_PYPDF2_AVAILABLE = False
PDF_PLUMBER_AVAILABLE = False
DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_PYPDF2_AVAILABLE = True
except ImportError:
    PyPDF2 = None

try:
    import pdfplumber  # type: ignore
    PDF_PLUMBER_AVAILABLE = True
except ImportError:
    pdfplumber = None

try:
    from docx import Document  # type: ignore
    DOCX_AVAILABLE = True
except ImportError:
    Document = None


def parse_markdown(file_path: str) -> tuple[str, dict]:
    """
    Parse markdown file and extract content and metadata.
    
    Args:
        file_path (str): Path to the markdown file
        
    Returns:
        tuple[str, dict]: Parsed content and metadata
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Extract basic metadata
        metadata = {
            'file_name': os.path.basename(file_path),
            'file_size': os.path.getsize(file_path),
            'file_type': 'markdown'
        }
        
        return content, metadata
    except Exception as e:
        logger.error(f"Error parsing markdown file {file_path}: {str(e)}", exc_info=True)
        raise


def parse_pdf(file_path: str) -> tuple[str, dict]:
    """
    Parse PDF file and extract content and metadata.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        tuple[str, dict]: Parsed content and metadata
    """
    content = ""
    metadata = {
        'file_name': os.path.basename(file_path),
        'file_size': os.path.getsize(file_path),
        'file_type': 'pdf'
    }
    
    # Try PyPDF2 first
    if PDF_PYPDF2_AVAILABLE and PyPDF2 is not None:
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                metadata['page_count'] = len(reader.pages)
                
                for page in reader.pages:
                    content += page.extract_text() + "\n"
            return content, metadata
        except Exception as e:
            logger.warning(f"Failed to parse PDF file with PyPDF2: {str(e)}")
            raise ValueError(f"Failed to parse PDF file with PyPDF2: {str(e)}")
    
    # Try pdfplumber as fallback
    elif PDF_PLUMBER_AVAILABLE and pdfplumber is not None:
        try:
            with open(file_path, 'rb') as f:
                with pdfplumber.open(f) as pdf:
                    metadata['page_count'] = len(pdf.pages)
                    
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            content += text + "\n"
            return content, metadata
        except Exception as e:
            logger.warning(f"Failed to parse PDF file with pdfplumber: {str(e)}")
            raise ValueError(f"Failed to parse PDF file with pdfplumber: {str(e)}")
    else:
        raise ImportError("No PDF parsing library available. Please install PyPDF2 or pdfplumber.")


def parse_docx(file_path: str) -> tuple[str, dict]:
    """
    Parse DOCX file and extract content and metadata.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        tuple[str, dict]: Parsed content and metadata
    """
    if not DOCX_AVAILABLE or Document is None:
        raise ImportError("DOCX parsing library not available. Please install python-docx.")
    
    try:
        doc = Document(file_path)
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Extract basic metadata
        metadata = {
            'file_name': os.path.basename(file_path),
            'file_size': os.path.getsize(file_path),
            'file_type': 'docx',
            'paragraph_count': len(doc.paragraphs)
        }
        
        return content, metadata
    except Exception as e:
        logger.error(f"Error parsing DOCX file {file_path}: {str(e)}", exc_info=True)
        raise ValueError(f"Failed to parse DOCX file: {str(e)}")


def ingest_document(file_path: str, file_type: str) -> tuple[str, dict]:
    """
    Ingest document of supported file types and extract content and metadata.
    
    Args:
        file_path (str): Path to the file
        file_type (str): Type of the file (markdown, pdf, docx)
        
    Returns:
        tuple[str, dict]: Parsed content and metadata
        
    Raises:
        ValueError: If file type is not supported or parsing fails
    """
    if file_type == 'markdown':
        return parse_markdown(file_path)
    elif file_type == 'pdf':
        return parse_pdf(file_path)
    elif file_type == 'docx':
        return parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def process_input_attachments(state: GlobalState) -> GlobalState:
    """
    Process input attachments and convert them to raw intelligence.
    
    Args:
        state (GlobalState): Current global state containing input attachments
        
    Returns:
        GlobalState: Updated state with processed attachments as raw intelligence
    """
    try:
        artifact_manager = ArtifactManager()
        
        for attachment in state.input_attachments:
            # Create a RawIntelligence entry for this attachment
            raw_intel = RawIntelligence(
                source_scout_id=f"attachment_{attachment.file_name}",
                content_summary=attachment.content_summary,
                artifact_ref_id=attachment.full_content_ref
            )
            state.raw_intelligence.append(raw_intel)
        
        return state
    except Exception as e:
        logger.error(f"Error in process_input_attachments: {str(e)}", exc_info=True)
        # Set error status in state
        state.status = GlobalStateStatus.AWAITING_USER
        # Add error info to state if there's a field for it, or re-raise
        raise


def intelligence_fusion(state: GlobalState) -> GlobalState:
    """
    Fuse and consolidate raw intelligence from scouts into a coherent briefing.
    
    This node acts as the Intelligence Officer, merging and deduplicating 
    information from multiple sources, resolving conflicts, and creating 
    a consolidated intelligence briefing for the CPSO.
    
    Args:
        state (GlobalState): Current global state containing raw intelligence
        
    Returns:
        GlobalState: Updated state with consolidated briefing
    """
    try:
        # Initialize LLM for IO
        llm = get_llm("cpso")  # Using CPSO level LLM for intelligence fusion
        
        # Initialize artifact manager
        artifact_manager = ArtifactManager()
        
        # Gather all raw intelligence content
        intelligence_contents = []
        for intelligence in state.raw_intelligence:
            # Retrieve artifact content
            artifact = artifact_manager.get_artifact(intelligence.artifact_ref_id)
            if artifact:
                content = f"Source Scout ID: {intelligence.source_scout_id}\n"
                content += f"Content Summary: {intelligence.content_summary}\n"
                content += f"Full Content: {artifact['content']}\n"
                content += "---\n"
                intelligence_contents.append(content)
        
        # If we have technical correction, include it with high priority
        technical_correction_content = ""
        if state.technical_correction:
            technical_correction_content = f"Technical Correction (High Priority):\n{state.technical_correction}\n---\n"
        
        # Create prompt for intelligence fusion
        prompt = f"""
You are the Intelligence Officer (IO).
Raw Intelligence Reports: 
{"".join(intelligence_contents)}

{technical_correction_content}
Task: Consolidate the raw intelligence into a single coherent briefing.

Instructions:
1. Merge and deduplicate information from multiple sources
2. If Technical Correction exists, OVERRIDE conflicting claims
3. Mark conflicting sources with [⚠️ CONFLICT]
4. Generate a Consolidated Intelligence Briefing (Markdown)

Requirements:
- Every fact must cite [Source: artifact_id]
- If data is missing, state "Data Missing: {{topic}}"
- Use clear section headings
- Focus on facts, not opinions

Output ONLY the consolidated briefing in Markdown format.
""".strip()
        
        # Generate consolidated briefing using LLM
        response = llm.invoke(prompt)
        consolidated_briefing = str(response.content)
        
        # Update state with consolidated briefing
        state.consolidated_briefing = consolidated_briefing
        state.status = GlobalStateStatus.DRAFTING
        
        return state
    except Exception as e:
        logger.error(f"Error in intelligence_fusion: {str(e)}", exc_info=True)
        # Set error status in state
        state.status = GlobalStateStatus.AWAITING_USER
        # Add error info to state if there's a field for it, or re-raise
        raise